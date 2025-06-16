import csv
import json
import re
from typing import List, Dict, Tuple
import anthropic

from unidecode import unidecode
from api_token import CLAUDE_API_KEY

ABSTRACT_EXPORT = 'Export_ESMRMB_2025_Abstract_20250520_141544.csv'
IMAGE_FOLDER = '/media/bigboy2/ESMRMB2025/image/'


client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

def create_parsing_prompt(text: str) -> str:
    """
    Create the prompt for Claude to parse the academic text.

    Args:
        text: The input text containing acknowledgments, data/code availability, and references

    Returns:
        Formatted prompt for Claude
    """
    prompt = f"""
Please parse the following academic text and extract three sections into a JSON structure. The text may contain acknowledgments, data and code availability statements, and references. Some sections might be missing or not clearly marked with headers.

Instructions:
1. Extract "Acknowledgments" - funding information, grants, institutional support
2. Extract "Data and Code Availability Statement" - information about data/code availability, access, repositories
3. Extract "References" as a list - each reference should be a separate item in the list
4. For references, remove any original numbering (like [1], 1., (1), etc.) and just include the reference text
5. If a section is not present, use null for that field
6. Do not modify or alter the original text - only organize it into the appropriate JSON fields
7. Infer section membership based on content when headers are unclear

Return ONLY valid JSON in this exact format:
{{
    "Acknowledgments": "text or null",
    "Data and Code Availability Statement": "text or null", 
    "References": ["reference1", "reference2", ...] or null
}}

Text to parse:
{text}
"""
    return prompt


def parse_refs(text: str):
    """
    Parse a single text using Claude API.

    Args:
        text: The academic text to parse

    Returns:
        Dictionary with parsed sections
    """
    try:
        prompt = create_parsing_prompt(text)

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            temperature=0,  # Use 0 for consistent parsing
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        response_text = message.content[0].text.strip()

        # Extract JSON from response (in case there's extra text)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
            return json.loads(json_str)
        else:
            # Try to parse the entire response as JSON
            return json.loads(response_text)

    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        print(f"Response was: {response_text}")
        return {
            "Acknowledgments": None,
            "Data and Code Availability Statement": None,
            "References": None,
            "error": f"JSON parsing failed: {str(e)}"
        }
    except Exception as e:
        print(f"API or other error: {e}")
        return {
            "Acknowledgments": None,
            "Data and Code Availability Statement": None,
            "References": None,
            "error": str(e)
        }


def parse_author_list(author_string: str) -> List[Tuple[str, str]]:
    """
    Parse an author list string into individual authors with their affiliations.

    Args:
        author_string: String containing authors with affiliations in parentheses

    Returns:
        List of tuples containing (author_name, affiliations)
    """
    authors = []

    # Pattern to match: Name (affiliations) followed by comma or end of string
    # This handles affiliations that may contain commas
    pattern = r'([^(,]+)\s*\(([^)]+)\)\s*(?:,\s*|$)'

    matches = re.findall(pattern, author_string)

    for name, affiliations in matches:
        # Clean up whitespace
        name = name.strip()
        affiliations = affiliations.strip()
        authors.append((name, affiliations))

    return authors


parentheses_re = re.compile(r'\([^)]*\)')

def remove_parentheses(line):
    return re.sub(parentheses_re, '', line)


# focus_topic, number, title, authors [list], keywords [list], general_audience_pitch [bool], text, primary_subcategory, secondary_subcategory

aff_cleanup_re = re.compile(r'^[0-9]+\.[\s,]*')
multispace_cleanup_re = re.compile(r'\s+')

fig_re = re.compile(r'^Fig(?:ure)?\.?\s*([0-9]+)\s*[:.]?\s*(.*)', re.IGNORECASE | re.MULTILINE)
table_re = re.compile(r'^Tab(?:le)?\.?\s*([0-9]+)\s*[:.]?\s*(.*)', re.IGNORECASE | re.MULTILINE)
def process_caption(caption: str):
    """
    Process a caption to extract figure/table number and description.

    Args:
        caption: The caption text to process

    Returns:
        Tuple containing (caption reference, cleaned caption)
    """
    fig_match = fig_re.match(caption)
    table_match = table_re.match(caption)
    if fig_match:
        return f'Figure {fig_match.group(1)} ', fig_match.group(2).strip()
    if table_match:
        return f'Table {table_match.group(1)} ', table_match.group(2).strip()

    print("Warning: Caption does not match expected format:", caption)
    return '', caption.strip()


def process_figure_field(field: str):
    """
    Process a field that may contain figure references.

    Args:
        field: The field to process

    Returns:
        Cleaned field with figure references removed
    """
    figure_file_re = re.compile(r'[0-9]{5}-[0-9]{6}-.*')
    figures = []
    caption_refs = []
    captions = []
    current_caption = None
    for line in field.splitlines():
        line = line.strip()
        if not line:
            continue
        if figure_file_re.match(line):
            if current_caption is not None:
                caption_ref, caption = process_caption(current_caption.strip())
                caption_refs.append(caption_ref)
                captions.append(caption)
            current_caption = ''
            figures.append(line)
        else:
            current_caption += line + '\n'
    if field and not figures:
        print("Error parsing figures in field:", field)
    if current_caption is not None:
        caption_ref, caption = process_caption(current_caption.strip())
        caption_refs.append(caption_ref)
        captions.append(caption)
    return figures, caption_refs, captions



with open(ABSTRACT_EXPORT, 'r', encoding='ISO-8859-15') as f:
    reader = csv.DictReader(f, delimiter=';', quotechar='"')
    abstracts = []
    row_number = 0
    for row in reader:
        #if row_number > 10:
        #    break
        row_number += 1

        print(f"Processing row {row_number}: {row['Reference']} - {row['Titre']}")

        if row['Statut'] != 'Reviewing Pending':
            continue
        abstract = {}
        abstract['title'] = unidecode(row['Titre'])
        abstract['reference'] = '#' + row['Reference']

        authors_line = re.sub(multispace_cleanup_re, ' ', unidecode(row['Auteurs']))
        authors_list = [unidecode(a.strip()) for a in authors_line.split(',')]
        abstract['authors'] = parse_author_list(authors_line)
        speaker = unidecode(row['Orateur nom'].strip())
        abstract['speaker'] = 0
        for i, (author_name, affiliations) in enumerate(abstract['authors']):
            if speaker.lower() in author_name.lower():
                abstract['speaker'] = i
                break

        affiliations = unidecode(row['Affiliations']).splitlines()
        # clean up affiliations
        abstract['affiliations'] = [re.sub(aff_cleanup_re, '', aff) for aff in affiliations if aff.strip()]

        abstract['introduction'] = unidecode(row['Résumé'])
        abstract['methods'] = unidecode(row['Methods'])
        abstract['results'] = unidecode(row['Results'])
        abstract['discussion'] = unidecode(row['Discussion'])
        abstract['conclusion'] = unidecode(row['Conclusion'])
        #refs = parse_refs(unidecode(row['Data and Code Availability Statement and References (Information not included in the word counting)']))
        refs = {'Acknowledgments': unidecode(row['Data and Code Availability Statement and References (Information not included in the word counting)'])}
        abstract['acknowledgments'] = refs.get('Acknowledgments', '')
        abstract['data_and_code_availability'] = refs.get('Data and Code Availability Statement', '')
        abstract['references'] = refs.get('References', [])
        abstract['figure_files'], abstract['figure_refs'], abstract['figure_captions'] = process_figure_field(unidecode(row['Figure']))
        abstracts.append(abstract)

        with open('abstracts_for_print.json', 'w', encoding='utf-8') as json_file:
            json.dump(abstracts, json_file, ensure_ascii=False, indent=4)

# create doc file
from docx import Document
from docx.shared import Inches
doc = Document()

ABSTRACT_SUBSET = """45995
46461
47569
47619
47669
47681
47826
47901""".splitlines()

for abstract in abstracts:
    if abstract['reference'][1:] not in ABSTRACT_SUBSET:
        continue
    title_par = doc.add_paragraph()
    title_par.add_run(abstract['reference'] + ' \n' + abstract['title']).bold = True
    authors_par = doc.add_paragraph('')
    for i, (author_name, affiliations) in enumerate(abstract['authors']):
        auth_run = authors_par.add_run(author_name)
        if i == abstract['speaker']:
            auth_run.underline = True
        aff_run = authors_par.add_run(affiliations)
        aff_run.font.superscript = True
        if (i + 1) < len(abstract['authors']):
            authors_par.add_run(', ')

    affiliations_par = doc.add_paragraph('')
    affiliations_par.italic = True
    for i, aff in enumerate(abstract['affiliations']):
        aff_run = affiliations_par.add_run(f'{i+1} {aff}').italic = True
        if (i + 1) < len(abstract['affiliations']):
            affiliations_par.add_run('\n')

    p = doc.add_paragraph('')
    p.add_run('Introduction: ').bold = True
    p.add_run(abstract['introduction'])

    p = doc.add_paragraph('')
    p.add_run('Methods: ').bold = True
    p.add_run(abstract['methods'])

    p = doc.add_paragraph('')
    p.add_run('Results: ').bold = True
    p.add_run(abstract['results'])

    p = doc.add_paragraph('')
    p.add_run('Discussion: ').bold = True
    p.add_run(abstract['discussion'])

    p = doc.add_paragraph('')
    p.add_run('Conclusion: ').bold = True
    p.add_run(abstract['conclusion'])

    if abstract['acknowledgments']:
        p = doc.add_paragraph('')
        p.add_run('Acknowledgments: ').bold = True
        p.add_run(abstract['acknowledgments'])

    if abstract['data_and_code_availability']:
        p = doc.add_paragraph('')
        p.add_run('Data and Code Availability: ').bold = True
        p.add_run(abstract['data_and_code_availability'])

    if abstract['figure_files']:
        for f_num, figure in enumerate(abstract['figure_files']):
            doc.add_picture(IMAGE_FOLDER + figure, width=Inches(5.0))
            p = doc.add_paragraph('')
            try:
                if not abstract['figure_refs'][f_num]:
                    p.add_run('Figure ' + str(f_num + 1) + ' ').bold = True
                else:
                    p.add_run(abstract['figure_refs'][f_num]).bold = True
                p.add_run(abstract['figure_captions'][f_num])
            except IndexError:
                print('Warning: Figure caption not found for figure', f_num + 1)

    if abstract['references']:
        p = doc.add_paragraph('')
        p.add_run('References:\n').bold = True
        for i, ref in enumerate(abstract['references']):
            ref_run = p.add_run(f'{i+1}. {ref}')
            if i+1 < len(abstract['references']):
                p.add_run('\n')


doc.save('abstracts_for_print.docx')

