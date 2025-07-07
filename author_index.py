import json

from docx import Document

with open('abstracts_merged.json', 'r', encoding='utf-8') as f:
    abstracts = json.load(f)

authors = {}

for abstract in abstracts:
    if not abstract['program_number']:
        continue
    for author in abstract['authors_separated']:
        first = author[0].strip()
        last = author[1].strip()
        first_names = first.split(' ')
        initials = ' '.join([f[0].upper() + '.' for f in first_names if f])
        name = (last, initials)
        if name not in authors:
            authors[name] = [abstract['program_number']]
        else:
            authors[name].append(abstract['program_number'])

doc = Document()
letter = None
for author in sorted(authors.keys(), key=lambda x: x[0].upper()):
    last, initials = author
    if last[0].upper() != letter:
        print(f'Adding new letter section for {last[0].upper()}')
        letter = last[0].upper()
        doc.add_heading(letter)
        par = doc.add_paragraph()


    par.add_run(f'{last}, {initials}')
    par.add_run('\t')
    par.add_run(', '.join([str(pn) for pn in sorted(authors[author], key=lambda x: int(x))]))
    par.add_run('\n')

doc.save('authors_index.docx')