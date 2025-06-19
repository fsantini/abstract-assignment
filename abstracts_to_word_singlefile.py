# create doc file
import pypandoc
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches
import json
import os
from unidecode import unidecode
import tempfile
from pdf2image import convert_from_path
from PIL import Image, ImageChops

IMAGE_FOLDER = '/media/bigboy2/ESMRMB2025/image/'
OUTPUT_FOLDER  = '/media/bigboy2/ESMRMB2025/abstracts/'

def crop_whitespace(img):
    bg = Image.new(img.mode, img.size, (255, 255, 255))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        return img.crop(bbox)
    return img

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

with open('abstracts_for_print.json', 'r', encoding='utf-8') as f:
    abstracts = json.load(f)

for i, abstract in enumerate(abstracts):
    print(f'Processing abstract {i+1}: {abstract["reference"]} - {abstract["title"]}')
    output_base_name = os.path.join(OUTPUT_FOLDER, f'{abstract['reference'][1:]}')
    if os.path.exists(output_base_name + '.pdf'):
        continue
    doc = Document()
    title_par = doc.add_paragraph()
    title_par.add_run(abstract['reference'] + ' \n' + abstract['title']).bold = True
    authors_par = doc.add_paragraph('')
    for i, (author_name, affiliations) in enumerate(abstract['authors']):
        auth_run = authors_par.add_run(author_name)
        if i == abstract['speaker']:
            auth_run.underline = True
        aff_run = authors_par.add_run(affiliations)
        aff_run.font.superscript = True
        if (i + 1) < len(abstract['authors']):
            authors_par.add_run(', ')

    affiliations_par = doc.add_paragraph('')
    affiliations_par.italic = True
    for i, aff in enumerate(abstract['affiliations']):
        aff_run = affiliations_par.add_run(f'{i+1} {aff}').italic = True
        if (i + 1) < len(abstract['affiliations']):
            affiliations_par.add_run('\n')

    p = doc.add_paragraph('')
    p.add_run('Introduction: ').bold = True
    p.add_run(abstract['introduction'])

    p = doc.add_paragraph('')
    p.add_run('Methods: ').bold = True
    p.add_run(abstract['methods'])

    p = doc.add_paragraph('')
    p.add_run('Results: ').bold = True
    p.add_run(abstract['results'])

    p = doc.add_paragraph('')
    p.add_run('Discussion: ').bold = True
    p.add_run(abstract['discussion'])

    p = doc.add_paragraph('')
    p.add_run('Conclusion: ').bold = True
    p.add_run(abstract['conclusion'])

    if abstract['acknowledgments']:
        p = doc.add_paragraph('')
        p.add_run('Acknowledgments: ').bold = True
        p.add_run(abstract['acknowledgments'])

    if abstract['data_and_code_availability']:
        p = doc.add_paragraph('')
        p.add_run('Data and Code Availability: ').bold = True
        p.add_run(abstract['data_and_code_availability'])

    if abstract['figure_files']:
        for f_num, figure in enumerate(abstract['figure_files']):
            figure = unidecode(figure.replace('?', ''))
            figure_path = os.path.join(IMAGE_FOLDER, figure)
            img_path = figure_path
            is_temp = False

            if figure.lower().endswith('.pdf'):
                print("Converting PDF figure to image:", figure_path)
                with tempfile.TemporaryDirectory() as tmpdir:
                    images = convert_from_path(figure_path, dpi=150, output_folder=tmpdir, fmt='png')
                    if images:
                        img = images[0]
                        img = crop_whitespace(img)
                        img_path = figure_path + '.png'
                        img.save(img_path)
                        is_temp = True
                    else:
                        print("Warning: PDF conversion failed for", figure_path)
                        continue
            try:
                doc.add_picture(img_path, width=Inches(5.0))
            except FileNotFoundError:
                print("Warning: Figure file not found:", IMAGE_FOLDER + figure)
                continue
            except UnrecognizedImageError:
                print("Warning: Unrecognized image format for file:", IMAGE_FOLDER + figure)
                continue
            p = doc.add_paragraph('')
            try:
                if not abstract['figure_refs'][f_num]:
                    p.add_run('Figure ' + str(f_num + 1) + ' ').bold = True
                else:
                    p.add_run(abstract['figure_refs'][f_num]).bold = True
                p.add_run(abstract['figure_captions'][f_num])
            except IndexError:
                print('Warning: Figure caption not found for figure', f_num + 1)

    if abstract['references']:
        p = doc.add_paragraph('')
        p.add_run('References:\n').bold = True
        for i, ref in enumerate(abstract['references']):
            ref_run = p.add_run(f'{i+1}. {ref}')
            if i+1 < len(abstract['references']):
                p.add_run('\n')


    output_base_name = os.path.join(OUTPUT_FOLDER, f'{abstract['reference'][1:]}')
    doc.save(output_base_name + '.docx')
    extra_args = [
        '-V', 'geometry:top=1.5in',
        '-V', 'geometry:bottom=1in',
        '-V', 'geometry:left=1.25in',
        '-V', 'geometry:right=1.25in'
    ]
    try:
        pypandoc.convert_file(output_base_name + '.docx', 'pdf', outputfile=output_base_name + '.pdf', extra_args=extra_args)
    except Exception as e:
        print(f"Error converting {output_base_name}.docx to PDF: {e}")
        continue

